import os

from PIL import Image
from docx import Document
from docx.shared import Inches

def add_to_word(
            filename,
            path_qr,
            qr_inches, 
            path_input, 
            path_output       
            ):
    document = Document(f"{path_input}\\{filename}.docx")
    section_odd = document.sections[0]
    header_odd = section_odd.footer
    paragraph_odd = header_odd.paragraphs[0]
    run_odd = paragraph_odd.add_run()
    run_odd.add_picture(path_qr, height = Inches(qr_inches[0])) #width = 75 , height = 75

    header_even = document.sections[0].even_page_footer
    paragraph_even = header_even.paragraphs[0]
    run_even = paragraph_even.add_run()
    run_even.add_picture(path_qr, height = Inches(qr_inches[1])) #width = 75 , height = 75

    document.save(f"{path_output}\\{filename}.docx")

    return f"{path_output}\\{filename}.docx"

def add_to_img(
            filename,
            path_qr, 
            qr_box,
            path_input,
            path_output
            ):
    qr = Image.open(path_qr).convert("RGB")
    
    pages = os.listdir(path_input)
    
    for page in pages:
        image = Image.open(f'{path_input}\\{page}')
        
        #Доделать
        if qr_box is None:
            left_indent = qr.height
            bottom_margin = -1
            for y in range(1, qr.width):         
                for x in range(1, qr.height):  
                    pixel = qr.getpixel((x, y))
                    if pixel[0] < 240 and pixel[1] < 240 and pixel[2] < 240:
                        if y > bottom_margin:
                            bottom_margin = y
                        elif x < left_indent:
                            left_indent = x 
            if left_indent != qr.height and bottom_margin != -1:
                qr_box = (left_indent + qr.width, image.height - bottom_margin )
            else:
                qr_box = (150, image.height - qr.height - 150 )

        image.paste(qr, qr_box)
        image = image.convert("RGB") 
        
        image.save(f'{path_output}\\{page}')
    
    return f'{path_output}'